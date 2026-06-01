#!/usr/bin/env python3
"""
file-info — Get metadata/info about any supported file

Usage:
  python3 scripts/file-info.py <path>

Supported: PDF, DOCX, XLSX, PPTX, images (JPG/PNG/GIF/BMP/TIFF/WEBP/ICO)

Output JSON:
  { path, filename, size_bytes, type, ... type-specific fields ... }
"""
import json
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from lib.detect import detect


def info_pdf(path):
    from pypdf import PdfReader
    r = PdfReader(path)
    meta = r.metadata or {}
    return {
        "pages": len(r.pages),
        "title": meta.get("/Title", ""),
        "author": meta.get("/Author", ""),
        "creator": meta.get("/Creator", ""),
        "encrypted": r.is_encrypted,
        "page_size": (
            {
                "width_pt": float(r.pages[0].mediabox.width),
                "height_pt": float(r.pages[0].mediabox.height),
            }
            if r.pages
            else None
        ),
    }


def info_docx(path):
    from docx import Document
    doc = Document(path)
    core = doc.core_properties
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "title": core.title or "",
        "author": core.author or "",
        "created": str(core.created) if core.created else "",
        "modified": str(core.modified) if core.modified else "",
        "word_count": sum(
            len(p.text.split()) for p in doc.paragraphs if p.text.strip()
        ),
    }


def info_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    sheets = []
    for name in wb.sheetnames:
        ws = wb[name]
        sheets.append({
            "name": name,
            "rows": ws.max_row,
            "columns": ws.max_column,
        })
    wb.close()
    return {"sheets": sheets, "sheet_count": len(sheets)}


def info_pptx(path):
    from pptx import Presentation
    from pptx.util import Emu
    prs = Presentation(path)
    w = prs.slide_width
    h = prs.slide_height
    return {
        "slides": len(prs.slides),
        "slide_width_px": round(w / 914400 * 96),   # EMU → px @96dpi
        "slide_height_px": round(h / 914400 * 96),
        "slide_layouts": len(prs.slide_layouts),
    }


def info_image(path):
    from PIL import Image, ExifTags
    img = Image.open(path)
    result = {
        "width": img.width,
        "height": img.height,
        "format": img.format or os.path.splitext(path)[1].lstrip(".").upper(),
        "mode": img.mode,
        "animated": getattr(img, "n_frames", 1) > 1,
        "frames": getattr(img, "n_frames", 1),
    }
    # EXIF if available
    try:
        exif_raw = img._getexif()
        if exif_raw:
            exif = {
                ExifTags.TAGS.get(k, k): str(v)
                for k, v in exif_raw.items()
                if k in ExifTags.TAGS
            }
            # Keep only common useful tags
            keep = {"Make", "Model", "DateTime", "GPSInfo", "Orientation",
                    "ExposureTime", "FNumber", "ISOSpeedRatings", "Software"}
            result["exif"] = {k: v for k, v in exif.items() if k in keep}
    except Exception:
        pass
    img.close()
    return result


def info_csv(path):
    with open(path, encoding="utf-8-sig", errors="replace") as f:
        lines = f.readlines()
    header = lines[0].strip() if lines else ""
    cols = len(header.split(",")) if header else 0
    return {
        "rows": len(lines) - 1,
        "columns": cols,
        "header": header[:200],
    }


def info_text(path):
    with open(path, encoding="utf-8-sig", errors="replace") as f:
        content = f.read()
    lines = content.splitlines()
    return {
        "lines": len(lines),
        "words": len(content.split()),
        "chars": len(content),
        "preview": content[:300].replace("\n", " "),
    }


HANDLERS = {
    "pdf":   info_pdf,
    "docx":  info_docx,
    "xlsx":  info_xlsx,
    "pptx":  info_pptx,
    "image": info_image,
    "csv":   info_csv,
    "tsv":   info_csv,
    "text":  info_text,
}


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: file-info.py <path>"}))
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.isfile(path):
        print(json.dumps({"error": f"File not found: {path}"}))
        sys.exit(1)

    ftype = detect(path)
    result = {
        "path":       os.path.abspath(path),
        "filename":   os.path.basename(path),
        "size_bytes": os.path.getsize(path),
        "type":       ftype,
    }

    handler = HANDLERS.get(ftype)
    if handler:
        try:
            result.update(handler(path))
        except Exception as e:
            result["error"] = str(e)
    else:
        result["note"] = f"No detailed info available for type '{ftype}'"

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
